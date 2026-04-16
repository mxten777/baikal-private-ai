"""
Document Loader - 텍스트 추출
"""
import logging
from app.config import get_settings

logger = logging.getLogger("baikal.loader")
settings = get_settings()


def extract_text(filepath: str, file_type: str) -> str:
    """파일에서 텍스트 추출"""
    extractors = {
        "pdf": extract_pdf,
        "docx": extract_docx,
        "xlsx": extract_xlsx,
        "hwp": extract_hwp,
        "hwpx": extract_hwpx,
    }

    extractor = extractors.get(file_type)
    if extractor is None:
        raise ValueError(f"지원하지 않는 파일 형식: {file_type}")

    try:
        text = extractor(filepath)
        logger.info(f"텍스트 추출 완료: {filepath} ({len(text)} chars)")
        return text
    except Exception as e:
        logger.error(f"텍스트 추출 실패: {filepath} - {e}")
        raise


def extract_pages(filepath: str, file_type: str) -> list:
    """파일에서 페이지/섹션별 텍스트 목록 반환.

    Returns: List[Tuple[Optional[int], str]] — (page_number, text)
    page_number는 1-based. None이면 페이지 추적 불가 형식.
    """
    try:
        if file_type == "pdf":
            return _extract_pdf_pages(filepath)
        elif file_type == "docx":
            return [(None, extract_docx(filepath))]
        elif file_type == "xlsx":
            return _extract_xlsx_pages(filepath)
        elif file_type == "hwp":
            return _extract_hwp_pages(filepath)
        elif file_type == "hwpx":
            return _extract_hwpx_pages(filepath)
        else:
            return [(None, extract_text(filepath, file_type))]
    except Exception as e:
        logger.warning(f"extract_pages 실패 ({file_type}): {e}")
        return []


def _extract_pdf_pages(filepath: str) -> list:
    """PDF에서 페이지별 (page_num, text) 반환"""
    try:
        import pdfplumber
    except ImportError:
        return [(None, extract_pdf(filepath))]

    pages = []
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages):
            try:
                page_text_parts = []
                tables = page.extract_tables()
                table_bboxes = [t.bbox for t in page.find_tables()] if tables else []

                for table in tables:
                    rows = []
                    for row in table:
                        cells = [str(cell).strip() if cell is not None else "" for cell in row]
                        rows.append("\t".join(cells))
                    page_text_parts.append("\n".join(rows))

                if table_bboxes:
                    plain = page.filter(
                        lambda obj: obj["object_type"] == "char" and not any(
                            obj["x0"] >= bbox[0] and obj["x1"] <= bbox[2]
                            and obj["top"] >= bbox[1] and obj["bottom"] <= bbox[3]
                            for bbox in table_bboxes
                        )
                    ).extract_text()
                else:
                    plain = page.extract_text()

                if plain and plain.strip():
                    page_text_parts.insert(0, plain.strip())

                if page_text_parts:
                    pages.append((i + 1, "\n".join(page_text_parts)))
            except Exception as e:
                logger.warning(f"PDF 페이지 {i + 1} 추출 실패: {e}")

    return pages if pages else [(None, extract_pdf(filepath))]


def _extract_xlsx_pages(filepath: str) -> list:
    """XLSX에서 시트별 (sheet_num, text) 반환"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return [(None, extract_xlsx(filepath))]

    wb = load_workbook(filepath, read_only=True, data_only=True)
    pages = []
    for i, sheet_name in enumerate(wb.sheetnames):
        ws = wb[sheet_name]
        text_parts = [f"[시트: {sheet_name}]"]
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
            if row_text.strip():
                text_parts.append(row_text)
                row_count += 1
            if row_count > 10000:
                break
        text = "\n".join(text_parts)
        if text.strip():
            pages.append((i + 1, text))
    wb.close()
    return pages if pages else [(None, extract_xlsx(filepath))]


def _extract_hwp_pages(filepath: str) -> list:
    """HWP에서 섹션별 (section_num, text) 반환"""
    try:
        import olefile
        import zlib
        import struct
    except ImportError:
        return [(None, extract_hwp(filepath))]

    if not olefile.isOleFile(filepath):
        return [(None, extract_hwp(filepath))]

    ole = olefile.OleFileIO(filepath)
    pages = []
    i = 1
    while True:
        stream_name = f'BodyText/Section{i:04d}'
        if not ole.exists(stream_name):
            break
        try:
            data = ole.openstream(stream_name).read()
            try:
                data = zlib.decompress(data, -15)
            except Exception:
                pass
            pos = 0
            text_parts = []
            while pos + 4 <= len(data):
                header = struct.unpack_from('<I', data, pos)[0]
                rec_type = header & 0x3FF
                size = (header >> 20) & 0xFFF
                if size == 0xFFF:
                    if pos + 8 > len(data):
                        break
                    size = struct.unpack_from('<I', data, pos + 4)[0]
                    pos += 8
                else:
                    pos += 4
                if rec_type == 67:
                    try:
                        text = data[pos:pos + size].decode('utf-16-le', errors='ignore')
                        text = ''.join(c for c in text if ord(c) >= 32 or c in '\n\t')
                        if text.strip():
                            text_parts.append(text)
                    except Exception:
                        pass
                pos += size
            if text_parts:
                pages.append((i, "\n".join(text_parts)))
        except Exception as e:
            logger.warning(f"HWP Section{i} 처리 실패: {e}")
        i += 1
    ole.close()
    return pages if pages else [(None, extract_hwp(filepath))]


def _extract_hwpx_pages(filepath: str) -> list:
    """HWPX에서 섹션 XML별 (section_num, text) 반환"""
    import zipfile
    import xml.etree.ElementTree as ET

    pages = []
    with zipfile.ZipFile(filepath, 'r') as z:
        section_files = sorted([
            f for f in z.namelist()
            if 'section' in f.lower() and f.endswith('.xml')
        ])
        for idx, section_file in enumerate(section_files):
            try:
                xml_data = z.read(section_file)
                root = ET.fromstring(xml_data)
                text_parts = []
                for elem in root.iter():
                    if elem.text and elem.text.strip():
                        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                        if tag in ('t', 'run', 'text', 'para'):
                            text_parts.append(elem.text.strip())
                if text_parts:
                    pages.append((idx + 1, "\n".join(text_parts)))
            except Exception as e:
                logger.warning(f"HWPX 섹션 파싱 실패 {section_file}: {e}")
    return pages if pages else []


def extract_pdf(filepath: str) -> str:
    """PDF에서 텍스트 추출 (표 구조 포함).
    pdfplumber로 텍스트 추출 시도 후, 추출된 텍스트가 너무 짧으면
    OCR(Tesseract)로 재시도.
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber가 설치되지 않았습니다: pip install pdfplumber")

    text_parts = []
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages):
            try:
                page_text_parts = []

                # 표 추출 (셀 단위로 탭 구분)
                tables = page.extract_tables()
                table_bboxes = [t.bbox for t in page.find_tables()] if tables else []

                for table in tables:
                    rows = []
                    for row in table:
                        cells = [str(cell).strip() if cell is not None else "" for cell in row]
                        rows.append("\t".join(cells))
                    page_text_parts.append("\n".join(rows))

                # 표 영역 제외한 일반 텍스트 추출
                if table_bboxes:
                    # 표가 있으면 표 제외 영역 텍스트만
                    plain = page.filter(
                        lambda obj: obj["object_type"] == "char" and not any(
                            obj["x0"] >= bbox[0] and obj["x1"] <= bbox[2]
                            and obj["top"] >= bbox[1] and obj["bottom"] <= bbox[3]
                            for bbox in table_bboxes
                        )
                    ).extract_text()
                else:
                    plain = page.extract_text()

                if plain and plain.strip():
                    page_text_parts.insert(0, plain.strip())

                if page_text_parts:
                    text_parts.append("\n".join(page_text_parts))

            except Exception as e:
                logger.warning(f"PDF 페이지 {i + 1} 추출 실패: {e}")
                continue

    result = "\n\n".join(text_parts)

    # 텍스트가 너무 짧으면 이미지 PDF로 판단 → OCR 시도
    with pdfplumber.open(filepath) as pdf:
        page_count = len(pdf.pages)

    if len(result.strip()) < settings.OCR_MIN_TEXT_PER_PAGE * max(page_count, 1):
        logger.info(f"텍스트 부족 ({len(result.strip())}자, {page_count}페이지) → OCR 시도: {filepath}")
        ocr_text = _extract_pdf_ocr(filepath)
        if ocr_text and len(ocr_text.strip()) > len(result.strip()):
            logger.info(f"OCR 텍스트 채택: {len(ocr_text)}자")
            return ocr_text

    if not result.strip():
        logger.warning(f"PDF에서 텍스트를 추출할 수 없습니다 (이미지 PDF일 수 있음): {filepath}")

    return result


def _extract_pdf_ocr(filepath: str) -> str:
    """Tesseract OCR로 PDF에서 텍스트 추출.
    pdf2image로 페이지를 이미지로 변환 후 pytesseract로 OCR.
    한국어(kor) + 영어(eng) 동시 인식.
    """
    try:
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError as e:
        logger.warning(f"OCR 패키지 미설치: {e}")
        return ""

    try:
        pages = convert_from_path(filepath, dpi=settings.OCR_DPI)
        text_parts = []
        for i, page_img in enumerate(pages):
            try:
                text = pytesseract.image_to_string(page_img, lang="kor+eng")
                if text.strip():
                    text_parts.append(text.strip())
            except Exception as e:
                logger.warning(f"OCR 페이지 {i + 1} 실패: {e}")
                continue
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.warning(f"OCR 처리 실패: {filepath} - {e}")
        return ""


def extract_docx(filepath: str) -> str:
    """DOCX에서 텍스트 추출"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx가 설치되지 않았습니다: pip install python-docx")

    doc = Document(filepath)
    text_parts = []

    # 본문 텍스트
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # 테이블 내용
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    # 헤더/푸터
    for section in doc.sections:
        header = section.header
        if header and header.paragraphs:
            for p in header.paragraphs:
                if p.text.strip():
                    text_parts.insert(0, p.text)

    return "\n".join(text_parts)


def extract_xlsx(filepath: str) -> str:
    """XLSX에서 텍스트 추출"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ImportError("openpyxl이 설치되지 않았습니다: pip install openpyxl")

    wb = load_workbook(filepath, read_only=True, data_only=True)
    text_parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        text_parts.append(f"[시트: {sheet_name}]")

        row_count = 0
        for row in ws.iter_rows(values_only=True):
            row_text = "\t".join(
                str(cell) if cell is not None else "" for cell in row
            )
            if row_text.strip():
                text_parts.append(row_text)
                row_count += 1

            # 대용량 시트 제한 (10000행)
            if row_count > 10000:
                text_parts.append(f"... ({sheet_name} 시트: 10000행까지만 처리)")
                break

    wb.close()
    return "\n".join(text_parts)


def extract_hwp(filepath: str) -> str:
    """HWP 파일에서 텍스트 추출 (OLE 바이너리 포맷)"""
    try:
        import olefile
        import zlib
        import struct
    except ImportError:
        raise ImportError("olefile이 설치되지 않았습니다: pip install olefile")

    if not olefile.isOleFile(filepath):
        raise ValueError("유효하지 않은 HWP 파일입니다")

    ole = olefile.OleFileIO(filepath)
    text_parts = []

    i = 1
    while True:
        stream_name = f'BodyText/Section{i:04d}'
        if not ole.exists(stream_name):
            break
        try:
            data = ole.openstream(stream_name).read()
            # zlib 압축 해제 시도
            try:
                data = zlib.decompress(data, -15)
            except Exception:
                pass

            # HWP 레코드 구조 파싱
            pos = 0
            while pos + 4 <= len(data):
                header = struct.unpack_from('<I', data, pos)[0]
                rec_type = header & 0x3FF
                size = (header >> 20) & 0xFFF
                if size == 0xFFF:
                    if pos + 8 > len(data):
                        break
                    size = struct.unpack_from('<I', data, pos + 4)[0]
                    pos += 8
                else:
                    pos += 4

                # HWPTAG_PARA_TEXT (rec_type 67 = 0x43)
                if rec_type == 67:
                    try:
                        text = data[pos:pos + size].decode('utf-16-le', errors='ignore')
                        text = ''.join(c for c in text if ord(c) >= 32 or c in '\n\t')
                        if text.strip():
                            text_parts.append(text)
                    except Exception:
                        pass
                pos += size
        except Exception as e:
            logger.warning(f"HWP Section{i} 처리 실패: {e}")
        i += 1

    ole.close()
    if not text_parts:
        logger.warning(f"HWP에서 텍스트를 추출하지 못했습니다: {filepath}")
    return "\n".join(text_parts)


def extract_hwpx(filepath: str) -> str:
    """HWPX 파일에서 텍스트 추출 (ZIP+XML 포맷)"""
    import zipfile
    import xml.etree.ElementTree as ET

    text_parts = []
    with zipfile.ZipFile(filepath, 'r') as z:
        section_files = sorted([
            f for f in z.namelist()
            if 'section' in f.lower() and f.endswith('.xml')
        ])
        for section_file in section_files:
            try:
                xml_data = z.read(section_file)
                root = ET.fromstring(xml_data)
                for elem in root.iter():
                    if elem.text and elem.text.strip():
                        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                        if tag in ('t', 'run', 'text', 'para'):
                            text_parts.append(elem.text.strip())
            except Exception as e:
                logger.warning(f"HWPX 섹션 파싱 실패 {section_file}: {e}")
    return "\n".join(text_parts)


# ---------------------------------------------------------------------------
# P3-8: 멀티모달 — qwen2.5-vl 이미지 설명 추출
# ---------------------------------------------------------------------------

_VISION_MODEL = "qwen2.5vl:7b"  # Ollama 비전 모델명


def _check_vision_model_available() -> bool:
    """qwen2.5-vl 모델이 Ollama에 설치됐는지 확인"""
    try:
        import httpx
        resp = httpx.get(f"{settings.OLLAMA_BASE_URL}/api/tags", timeout=5.0)
        if resp.status_code == 200:
            models = [m["name"] for m in resp.json().get("models", [])]
            return any("qwen2.5vl" in m.lower() or "qwen2.5-vl" in m.lower() for m in models)
    except Exception:
        pass
    return False


def describe_image(image_bytes: bytes, prompt: str | None = None) -> str:
    """이미지를 qwen2.5-vl로 분석하여 텍스트 설명 반환 (P3-8).

    qwen2.5-vl 미설치 시 빈 문자열 반환 (graceful fallback).
    """
    if not _check_vision_model_available():
        logger.debug("qwen2.5-vl 미설치 — 이미지 설명 건너뜀")
        return ""

    import base64
    import httpx

    if prompt is None:
        prompt = (
            "이 이미지에서 차트, 표, 그림, 다이어그램의 내용을 상세하게 한국어로 설명하세요. "
            "숫자, 제목, 범례, 축 레이블 등 모든 텍스트 정보를 포함해주세요."
        )

    b64_image = base64.b64encode(image_bytes).decode("utf-8")

    payload = {
        "model": _VISION_MODEL,
        "prompt": prompt,
        "images": [b64_image],
        "stream": False,
        "options": {"temperature": 0.1},
    }

    try:
        resp = httpx.post(
            f"{settings.OLLAMA_BASE_URL}/api/generate",
            json=payload,
            timeout=120.0,
        )
        resp.raise_for_status()
        description = resp.json().get("response", "").strip()
        logger.info(f"이미지 설명 추출 완료 ({len(description)}자)")
        return description
    except Exception as e:
        logger.warning(f"이미지 설명 추출 실패: {e}")
        return ""


def extract_pdf_with_vision(filepath: str) -> str:
    """PDF 텍스트 추출 + 이미지 페이지는 qwen2.5-vl로 분석 (P3-8).

    텍스트가 충분한 페이지: 기존 pdfplumber 추출 유지
    이미지 중심 페이지 (텍스트 < OCR_MIN_TEXT_PER_PAGE): qwen2.5-vl로 분석
    qwen2.5-vl 미설치 시: 기존 extract_pdf() 동작과 동일
    """
    if not _check_vision_model_available():
        return extract_pdf(filepath)

    try:
        import pdfplumber
        from pdf2image import convert_from_path
    except ImportError:
        return extract_pdf(filepath)

    min_chars = settings.OCR_MIN_TEXT_PER_PAGE
    text_parts = []

    with pdfplumber.open(filepath) as pdf:
        # 이미지 변환 (전 페이지)
        try:
            page_images = convert_from_path(filepath, dpi=150)
        except Exception:
            page_images = []

        for i, page in enumerate(pdf.pages):
            page_num = i + 1
            try:
                # 기존 텍스트 추출
                plain = page.extract_text() or ""
                tables = page.extract_tables()
                table_text_parts = []
                for table in tables:
                    rows = []
                    for row in table:
                        cells = [str(c).strip() if c else "" for c in row]
                        rows.append("\t".join(cells))
                    table_text_parts.append("\n".join(rows))

                combined = "\n".join(
                    ([plain.strip()] if plain.strip() else []) + table_text_parts
                ).strip()

                if len(combined) >= min_chars:
                    # 텍스트 충분 — 기존 방식 유지
                    text_parts.append(combined)
                elif i < len(page_images):
                    # 이미지 페이지 — qwen2.5-vl 분석
                    logger.info(f"PDF 페이지 {page_num}: 텍스트 부족 ({len(combined)}자) → 비전 모델 분석")
                    import io
                    img_buf = io.BytesIO()
                    page_images[i].save(img_buf, format="PNG")
                    description = describe_image(img_buf.getvalue())
                    if description:
                        text_parts.append(f"[페이지 {page_num} 이미지 내용]\n{description}")
                    elif combined:
                        text_parts.append(combined)
                else:
                    if combined:
                        text_parts.append(combined)

            except Exception as e:
                logger.warning(f"PDF 페이지 {page_num} 비전 처리 실패: {e}")

    return "\n\n".join(text_parts)


def extract_docx_with_vision(filepath: str) -> str:
    """DOCX 텍스트 추출 + 삽입 이미지는 qwen2.5-vl로 분석 (P3-8).

    qwen2.5-vl 미설치 시: 기존 extract_docx() 동작과 동일
    """
    if not _check_vision_model_available():
        return extract_docx(filepath)

    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        return extract_docx(filepath)

    doc = Document(filepath)
    text_parts = []
    image_count = 0

    # 인라인 이미지 추출 헬퍼
    def _get_inline_images(part):
        """docx 파트에서 인라인 이미지 바이트 리스트 반환"""
        images = []
        try:
            rels = part.rels
            for rel in rels.values():
                if "image" in rel.reltype:
                    try:
                        images.append(rel.target_part.blob)
                    except Exception:
                        pass
        except Exception:
            pass
        return images

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
        # 단락 내 이미지 처리
        for run in para.runs:
            drawing_elems = run._element.findall('.//' + qn('a:blip'))
            if drawing_elems:
                for blip in drawing_elems:
                    embed = blip.get(qn('r:embed'))
                    if embed and embed in doc.part.rels:
                        try:
                            img_blob = doc.part.rels[embed].target_part.blob
                            description = describe_image(img_blob)
                            if description:
                                image_count += 1
                                text_parts.append(f"[이미지 {image_count} 내용]\n{description}")
                        except Exception as e:
                            logger.warning(f"DOCX 이미지 처리 실패: {e}")

    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                text_parts.append(row_text)

    if image_count > 0:
        logger.info(f"DOCX 이미지 {image_count}개 비전 분석 완료")

    return "\n".join(text_parts)

