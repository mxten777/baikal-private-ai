"""
BAIKAL Private AI - 시연용 샘플 문서 생성 스크립트
실행: python scripts/create_demo_docs.py
"""
import os
from pathlib import Path

# 출력 디렉토리
OUTPUT_DIR = Path(__file__).parent.parent / "demo_docs"
OUTPUT_DIR.mkdir(exist_ok=True)


# ============================================================
# 1. PDF - 회사 취업규칙 (사내 규정)
# ============================================================
def create_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.units import mm

    filepath = OUTPUT_DIR / "바이칼_취업규칙.pdf"

    # 한글 폰트 등록 (Windows 기본 맑은고딕)
    font_path = "C:/Windows/Fonts/malgun.ttf"
    if os.path.exists(font_path):
        pdfmetrics.registerFont(TTFont("Malgun", font_path))
        font_name = "Malgun"
    else:
        font_name = "Helvetica"

    doc = SimpleDocTemplate(str(filepath), pagesize=A4,
                            topMargin=25*mm, bottomMargin=25*mm)

    title_style = ParagraphStyle("Title", fontName=font_name, fontSize=18,
                                  leading=28, alignment=1, spaceAfter=20)
    h1_style = ParagraphStyle("H1", fontName=font_name, fontSize=14,
                               leading=22, spaceBefore=16, spaceAfter=8)
    body_style = ParagraphStyle("Body", fontName=font_name, fontSize=11,
                                 leading=18, spaceAfter=6)

    content = []

    content.append(Paragraph("(주)바이칼테크놀로지 취업규칙", title_style))
    content.append(Spacer(1, 12))

    sections = [
        ("제1장 총칙", [
            "제1조 (목적) 이 규칙은 (주)바이칼테크놀로지(이하 '회사')의 직원 채용, 복무, 급여, 퇴직 등에 관한 사항을 규정함을 목적으로 한다.",
            "제2조 (적용범위) 이 규칙은 회사에 근무하는 모든 직원에게 적용된다. 다만, 계약직·파견직 등 특수 고용 형태의 직원에 대해서는 별도 계약에 따른다.",
            "제3조 (용어의 정의) ① '직원'이란 이 규칙에 따라 채용된 자를 말한다. ② '정규직'이란 수습기간을 마치고 정식 임용된 직원을 말한다.",
        ]),
        ("제2장 채용 및 수습", [
            "제4조 (채용) ① 회사는 서류전형, 면접전형을 거쳐 직원을 채용한다. ② 채용 시 이력서, 자기소개서, 졸업증명서, 경력증명서 등을 제출하여야 한다.",
            "제5조 (수습기간) ① 신규 입사자의 수습기간은 3개월로 한다. ② 수습기간 중 업무 적격 여부를 평가하여 정식 임용 또는 채용 취소할 수 있다.",
            "제6조 (결격사유) 다음 각 호에 해당하는 자는 직원으로 채용할 수 없다. 1. 금치산자 또는 한정치산자 2. 파산선고를 받은 자로서 복권되지 아니한 자",
        ]),
        ("제3장 근무시간 및 휴가", [
            "제7조 (근무시간) ① 1일 근무시간은 8시간, 주당 40시간을 원칙으로 한다. ② 근무시간은 오전 9시부터 오후 6시까지로 하며, 점심시간은 12시부터 1시까지로 한다.",
            "제8조 (유연근무제) ① 회사는 업무 특성에 따라 시차출퇴근제, 재택근무제를 시행할 수 있다. ② 시차출퇴근제의 경우 오전 8시~10시 사이 출근, 8시간 근무 후 퇴근한다.",
            "제9조 (연차휴가) ① 1년간 80% 이상 출근한 직원에게 15일의 연차유급휴가를 부여한다. ② 입사 1년 미만 직원에게는 1개월 개근 시 1일의 유급휴가를 부여한다. ③ 3년 이상 근속 시 매 2년마다 1일을 가산하되, 총 25일을 초과할 수 없다.",
            "제10조 (경조사휴가) 다음 경조사 발생 시 유급휴가를 부여한다: 본인 결혼 5일, 자녀 결혼 2일, 부모 사망 5일, 배우자 사망 5일, 조부모/형제자매 사망 3일, 배우자 출산 10일.",
            "제11조 (병가) ① 업무 외 부상 또는 질병으로 근무가 불가능한 경우 연간 60일 이내의 병가를 신청할 수 있다. ② 병가 3일 이상 시 의사 진단서를 제출하여야 한다.",
        ]),
        ("제4장 급여 및 복리후생", [
            "제12조 (급여) ① 급여는 기본급, 직무수당, 성과급으로 구성된다. ② 급여 지급일은 매월 25일로 하며, 지급일이 휴일인 경우 전일에 지급한다.",
            "제13조 (초과근무수당) ① 1일 8시간, 주 40시간을 초과하는 근무에 대해 통상임금의 150%를 지급한다. ② 야간근무(22시~06시)에 대해 통상임금의 50%를 가산 지급한다.",
            "제14조 (복리후생) 회사는 다음의 복리후생 제도를 운영한다: ① 4대 보험 가입 ② 경조금 지급 (결혼 100만원, 출산 50만원, 사망 100만원) ③ 자기개발비 연 200만원 ④ 건강검진 연 1회 ⑤ 중식비 월 15만원 ⑥ 통신비 월 5만원",
            "제15조 (퇴직급여) 1년 이상 근속한 직원이 퇴직하는 경우, 근속년수 1년에 대해 30일분의 평균임금을 퇴직금으로 지급한다.",
        ]),
        ("제5장 보안 및 비밀유지", [
            "제16조 (비밀유지의무) ① 직원은 재직 중은 물론 퇴직 후에도 업무상 알게 된 회사의 기밀을 누설하여서는 안 된다. ② 기밀의 범위: 기술정보, 영업비밀, 고객정보, 내부 경영정보 등",
            "제17조 (정보보안) ① 회사의 정보시스템은 업무 목적으로만 사용하여야 한다. ② 개인용 USB, 외부 클라우드 서비스에 업무 자료를 저장하는 것을 금지한다. ③ 사내 AI 시스템(BAIKAL Private AI)을 통한 문서 검색은 허용하나, 외부 AI 서비스(ChatGPT 등) 사용은 금지한다.",
            "제18조 (위반 시 조치) 보안 규정 위반 시 경고, 감봉, 정직, 해고 등의 징계 조치를 할 수 있다.",
        ]),
    ]

    for title, articles in sections:
        content.append(Paragraph(title, h1_style))
        for article in articles:
            content.append(Paragraph(article, body_style))

    content.append(Spacer(1, 20))
    content.append(Paragraph("부칙: 이 규칙은 2025년 1월 1일부터 시행한다.", body_style))
    content.append(Paragraph("(주)바이칼테크놀로지 대표이사", body_style))

    doc.build(content)
    print(f"  ✅ PDF 생성: {filepath}")
    return filepath


# ============================================================
# 2. PDF (reportlab 없을 때 대안) - PyPDF2는 읽기 전용이라 fpdf 사용
# ============================================================
def create_pdf_simple():
    """reportlab 없을 때 텍스트 파일로 대체"""
    filepath = OUTPUT_DIR / "바이칼_취업규칙.txt"
    text = """(주)바이칼테크놀로지 취업규칙

제1장 총칙
제1조 (목적) 이 규칙은 (주)바이칼테크놀로지(이하 '회사')의 직원 채용, 복무, 급여, 퇴직 등에 관한 사항을 규정함을 목적으로 한다.
제2조 (적용범위) 이 규칙은 회사에 근무하는 모든 직원에게 적용된다.
제3조 (용어의 정의) ① '직원'이란 이 규칙에 따라 채용된 자를 말한다. ② '정규직'이란 수습기간을 마치고 정식 임용된 직원을 말한다.

제2장 채용 및 수습
제4조 (채용) 회사는 서류전형, 면접전형을 거쳐 직원을 채용한다.
제5조 (수습기간) 신규 입사자의 수습기간은 3개월로 한다.

제3장 근무시간 및 휴가
제7조 (근무시간) 1일 근무시간은 8시간, 주당 40시간을 원칙으로 한다. 근무시간은 오전 9시부터 오후 6시까지로 한다.
제8조 (유연근무제) 시차출퇴근제의 경우 오전 8시~10시 사이 출근, 8시간 근무 후 퇴근한다.
제9조 (연차휴가) 1년간 80% 이상 출근한 직원에게 15일의 연차유급휴가를 부여한다. 3년 이상 근속 시 매 2년마다 1일을 가산하되, 총 25일을 초과할 수 없다.
제10조 (경조사휴가) 본인 결혼 5일, 자녀 결혼 2일, 부모 사망 5일, 배우자 출산 10일.
제11조 (병가) 업무 외 부상 또는 질병으로 근무가 불가능한 경우 연간 60일 이내의 병가를 신청할 수 있다.

제4장 급여 및 복리후생
제12조 (급여) 급여 지급일은 매월 25일로 한다.
제13조 (초과근무수당) 1일 8시간을 초과하는 근무에 대해 통상임금의 150%를 지급한다. 야간근무(22시~06시)는 50% 가산.
제14조 (복리후생) 경조금(결혼 100만원, 출산 50만원), 자기개발비 연 200만원, 건강검진 연 1회, 중식비 월 15만원, 통신비 월 5만원.
제15조 (퇴직급여) 근속년수 1년에 대해 30일분의 평균임금을 퇴직금으로 지급한다.

제5장 보안 및 비밀유지
제16조 (비밀유지의무) 직원은 재직 중은 물론 퇴직 후에도 업무상 알게 된 회사의 기밀을 누설하여서는 안 된다.
제17조 (정보보안) 개인용 USB, 외부 클라우드 서비스에 업무 자료를 저장하는 것을 금지한다. 사내 AI 시스템(BAIKAL Private AI)을 통한 문서 검색은 허용하나, 외부 AI 서비스 사용은 금지한다.

부칙: 이 규칙은 2025년 1월 1일부터 시행한다.
"""
    filepath.write_text(text, encoding="utf-8")
    print(f"  ✅ TXT 생성 (PDF 대체): {filepath}")
    return filepath


# ============================================================
# 3. DOCX - 제품 사용 매뉴얼
# ============================================================
def create_docx():
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    filepath = OUTPUT_DIR / "BAIKAL_AI_사용매뉴얼.docx"
    doc = Document()

    # 제목
    title = doc.add_heading("BAIKAL Private AI 사용 매뉴얼", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("버전 1.0 | 2025년 1월 | (주)바이칼테크놀로지")
    doc.add_paragraph("")

    # 1. 개요
    doc.add_heading("1. 시스템 개요", level=1)
    doc.add_paragraph(
        "BAIKAL Private AI는 폐쇄망(에어갭) 환경에서 운영 가능한 문서 기반 AI 질의응답 플랫폼입니다. "
        "모든 데이터 처리가 사내 서버에서 이루어지므로 외부 유출 위험이 없습니다."
    )
    doc.add_paragraph("주요 특징:", style="List Bullet")
    for feature in [
        "외부 API 의존 없음 — 완전한 오프라인 AI",
        "문서 업로드 → 자동 분석 → 벡터 검색 → AI 답변",
        "실시간 스트리밍 응답 (ChatGPT와 유사한 UX)",
        "PDF, DOCX, XLSX 문서 형식 지원",
        "관리자/일반 사용자 역할 기반 접근 제어",
    ]:
        doc.add_paragraph(feature, style="List Bullet")

    # 2. 로그인
    doc.add_heading("2. 로그인 방법", level=1)
    doc.add_paragraph(
        "웹 브라우저에서 시스템 주소(예: http://서버IP)에 접속합니다. "
        "관리자로부터 부여받은 아이디와 비밀번호를 입력하고 '로그인' 버튼을 클릭합니다."
    )
    doc.add_paragraph("초기 관리자 계정: admin / admin1234 (최초 로그인 후 반드시 비밀번호 변경)")

    # 3. 문서 업로드
    doc.add_heading("3. 문서 업로드", level=1)
    doc.add_paragraph(
        "좌측 메뉴에서 '문서 관리'를 클릭합니다. '파일 업로드' 영역에 문서를 드래그 앤 드롭하거나, "
        "클릭하여 파일을 선택합니다."
    )
    table = doc.add_table(rows=4, cols=3, style="Table Grid")
    headers = ["파일 형식", "확장자", "최대 크기"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ("PDF 문서", ".pdf", "100MB"),
        ("Word 문서", ".docx", "100MB"),
        ("Excel 문서", ".xlsx", "100MB"),
    ]
    for r, (a, b, c) in enumerate(data, 1):
        table.rows[r].cells[0].text = a
        table.rows[r].cells[1].text = b
        table.rows[r].cells[2].text = c

    doc.add_paragraph("")
    doc.add_paragraph(
        "업로드된 문서는 자동으로 텍스트 추출 → 청킹 → 벡터 임베딩 과정을 거칩니다. "
        "처리 상태는 '처리중' → '완료'로 변경됩니다. 완료 후 AI 질문응답에서 해당 문서를 활용할 수 있습니다."
    )

    # 4. AI 질문응답
    doc.add_heading("4. AI 질문응답 (채팅)", level=1)
    doc.add_paragraph(
        "좌측 메뉴에서 'AI 질문응답'을 클릭합니다. '새 대화' 버튼으로 대화 세션을 생성하고, "
        "하단 입력창에 질문을 입력합니다."
    )
    doc.add_heading("질문 예시", level=2)
    for q in [
        "연차휴가는 몇 일인가요?",
        "초과근무수당은 어떻게 계산되나요?",
        "이 문서를 요약해주세요",
        "복리후생 제도에는 어떤 것들이 있나요?",
        "신입사원 수습기간은 얼마인가요?",
    ]:
        doc.add_paragraph(q, style="List Bullet")

    doc.add_paragraph(
        "AI는 업로드된 문서에서 관련 내용을 찾아 답변합니다. "
        "답변 하단에 참고한 문서와 해당 부분이 표시됩니다."
    )

    # 5. 문서 검색
    doc.add_heading("5. 문서 검색", level=1)
    doc.add_paragraph(
        "'문서 검색' 메뉴에서 키워드를 입력하면 의미 기반(시맨틱) 검색을 수행합니다. "
        "단순 키워드 매칭이 아닌, 문맥과 의미를 파악하여 관련 문서 구간을 찾아줍니다. "
        "각 결과에는 유사도 점수가 함께 표시됩니다."
    )

    # 6. 관리자 기능
    doc.add_heading("6. 관리자 기능", level=1)
    doc.add_paragraph("관리자(admin) 역할로 로그인하면 추가 메뉴가 표시됩니다:")
    for item in [
        "사용자 관리: 사용자 생성/삭제, 역할(admin/user) 변경, 활성/비활성 전환",
        "문서 관리(관리자): 전체 사용자 문서 조회, 문서 삭제",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # 7. FAQ
    doc.add_heading("7. 자주 묻는 질문 (FAQ)", level=1)
    faqs = [
        ("Q: AI 답변의 정확도는?", "A: AI는 업로드된 문서에 기반하여 답변합니다. 문서에 없는 내용은 답변하지 못할 수 있습니다. 구체적인 질문일수록 정확한 답변을 얻을 수 있습니다."),
        ("Q: 동시 접속자 수 제한은?", "A: 서버 사양에 따라 다르지만, 일반적으로 20~50명 동시 접속을 지원합니다."),
        ("Q: 지원하지 않는 파일 형식은?", "A: 현재 PDF, DOCX, XLSX만 지원합니다. HWP, 이미지 파일은 향후 지원 예정입니다."),
        ("Q: 데이터는 외부로 전송되나요?", "A: 아니오. 모든 처리가 사내 서버에서 이루어지며, 외부 네트워크와 완전히 분리되어 있습니다."),
    ]
    for q, a in faqs:
        doc.add_paragraph(q).bold = True
        doc.add_paragraph(a)

    doc.save(str(filepath))
    print(f"  ✅ DOCX 생성: {filepath}")
    return filepath


# ============================================================
# 4. XLSX - 월별 매출 현황
# ============================================================
def create_xlsx():
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter

    filepath = OUTPUT_DIR / "2025년_매출현황.xlsx"
    wb = Workbook()

    # --- Sheet 1: 월별 매출 ---
    ws1 = wb.active
    ws1.title = "월별매출"

    header_fill = PatternFill(start_color="4F46E5", end_color="4F46E5", fill_type="solid")
    header_font = Font(name="맑은 고딕", size=11, bold=True, color="FFFFFF")
    data_font = Font(name="맑은 고딕", size=10)
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    # 제목
    ws1.merge_cells("A1:F1")
    ws1["A1"] = "(주)바이칼테크놀로지 2025년 월별 매출 현황"
    ws1["A1"].font = Font(name="맑은 고딕", size=14, bold=True)
    ws1["A1"].alignment = Alignment(horizontal="center")

    # 헤더
    headers = ["월", "솔루션 매출", "SI 매출", "유지보수 매출", "합계", "전년 대비(%)"]
    for col, h in enumerate(headers, 1):
        cell = ws1.cell(row=3, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin_border

    # 데이터
    monthly_data = [
        ("1월",  1200, 800, 350, None, 112),
        ("2월",  1350, 750, 350, None, 108),
        ("3월",  1800, 1200, 360, None, 125),
        ("4월",  1500, 900, 360, None, 115),
        ("5월",  1650, 1100, 370, None, 118),
        ("6월",  2100, 1500, 370, None, 132),
        ("7월",  1900, 1300, 380, None, 122),
        ("8월",  1400, 850, 380, None, 105),
        ("9월",  2200, 1600, 390, None, 135),
        ("10월", 2500, 1800, 390, None, 142),
        ("11월", 2800, 2000, 400, None, 148),
        ("12월", 3200, 2500, 400, None, 155),
    ]

    for r, (month, sol, si, maint, _, yoy) in enumerate(monthly_data, 4):
        ws1.cell(row=r, column=1, value=month).font = data_font
        ws1.cell(row=r, column=2, value=sol).font = data_font
        ws1.cell(row=r, column=3, value=si).font = data_font
        ws1.cell(row=r, column=4, value=maint).font = data_font
        ws1.cell(row=r, column=5).value = sol + si + maint
        ws1.cell(row=r, column=5).font = Font(name="맑은 고딕", size=10, bold=True)
        ws1.cell(row=r, column=6, value=yoy).font = data_font
        ws1.cell(row=r, column=6).number_format = '0"%"'
        for c in range(1, 7):
            ws1.cell(row=r, column=c).border = thin_border
            ws1.cell(row=r, column=c).alignment = Alignment(horizontal="center")
            if c in (2, 3, 4, 5):
                ws1.cell(row=r, column=c).number_format = '#,##0"만원"'

    # 합계 행
    sum_row = 16
    ws1.cell(row=sum_row, column=1, value="합계").font = Font(name="맑은 고딕", size=10, bold=True)
    for c in range(2, 6):
        ws1.cell(row=sum_row, column=c).value = sum(
            ws1.cell(row=r, column=c).value for r in range(4, 16)
        )
        ws1.cell(row=sum_row, column=c).font = Font(name="맑은 고딕", size=10, bold=True)
        ws1.cell(row=sum_row, column=c).number_format = '#,##0"만원"'
        ws1.cell(row=sum_row, column=c).border = thin_border
    ws1.cell(row=sum_row, column=1).border = thin_border
    ws1.cell(row=sum_row, column=6).border = thin_border

    # 열 너비
    widths = [8, 16, 14, 16, 14, 14]
    for i, w in enumerate(widths, 1):
        ws1.column_dimensions[get_column_letter(i)].width = w

    # --- Sheet 2: 제품별 매출 ---
    ws2 = wb.create_sheet("제품별매출")
    ws2.merge_cells("A1:D1")
    ws2["A1"] = "2025년 제품별 매출 분석"
    ws2["A1"].font = Font(name="맑은 고딕", size=14, bold=True)

    headers2 = ["제품명", "매출(만원)", "비율(%)", "주요 고객"]
    for col, h in enumerate(headers2, 1):
        cell = ws2.cell(row=3, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    products = [
        ("BAIKAL Private AI", 12500, 35.2, "삼성전자, LG전자, 국방부"),
        ("BAIKAL Document Hub", 8200, 23.1, "현대자동차, SK하이닉스"),
        ("BAIKAL Analytics", 6800, 19.2, "포스코, 한화, KT"),
        ("BAIKAL Security Suite", 4500, 12.7, "금융감독원, 국민은행"),
        ("기타 SI/컨설팅", 3500, 9.8, "다수"),
    ]
    for r, (name, revenue, ratio, clients) in enumerate(products, 4):
        ws2.cell(row=r, column=1, value=name).font = data_font
        ws2.cell(row=r, column=2, value=revenue).font = data_font
        ws2.cell(row=r, column=2).number_format = '#,##0'
        ws2.cell(row=r, column=3, value=ratio).font = data_font
        ws2.cell(row=r, column=4, value=clients).font = data_font

    for i, w in enumerate([22, 14, 10, 30], 1):
        ws2.column_dimensions[get_column_letter(i)].width = w

    # --- Sheet 3: 인력 현황 ---
    ws3 = wb.create_sheet("인력현황")
    ws3.merge_cells("A1:D1")
    ws3["A1"] = "2025년 부서별 인력 현황"
    ws3["A1"].font = Font(name="맑은 고딕", size=14, bold=True)

    headers3 = ["부서", "인원수", "평균 연차(년)", "비고"]
    for col, h in enumerate(headers3, 1):
        cell = ws3.cell(row=3, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill

    depts = [
        ("AI연구소", 25, 4.2, "머신러닝, NLP 전문"),
        ("개발1팀", 18, 5.1, "백엔드, 인프라"),
        ("개발2팀", 15, 3.8, "프론트엔드, 모바일"),
        ("영업팀", 12, 6.3, "공공, 기업 영업"),
        ("기획팀", 8, 4.5, "제품 기획, UX"),
        ("경영지원", 7, 7.2, "인사, 재무, 총무"),
        ("보안팀", 5, 8.1, "정보보안, 컴플라이언스"),
    ]
    for r, (dept, cnt, avg_yr, note) in enumerate(depts, 4):
        ws3.cell(row=r, column=1, value=dept).font = data_font
        ws3.cell(row=r, column=2, value=cnt).font = data_font
        ws3.cell(row=r, column=3, value=avg_yr).font = data_font
        ws3.cell(row=r, column=4, value=note).font = data_font

    ws3.cell(row=11, column=1, value="합계").font = Font(name="맑은 고딕", bold=True)
    ws3.cell(row=11, column=2, value=sum(d[1] for d in depts)).font = Font(name="맑은 고딕", bold=True)

    for i, w in enumerate([14, 10, 14, 24], 1):
        ws3.column_dimensions[get_column_letter(i)].width = w

    wb.save(str(filepath))
    print(f"  ✅ XLSX 생성: {filepath}")
    return filepath


# ============================================================
# 메인 실행
# ============================================================
if __name__ == "__main__":
    print()
    print("=" * 50)
    print(" BAIKAL Private AI - 시연용 샘플 문서 생성")
    print("=" * 50)
    print(f" 저장 위치: {OUTPUT_DIR}")
    print()

    # PDF 생성
    try:
        create_pdf()
    except ImportError:
        print("  ⚠️  reportlab 미설치 → 텍스트 파일로 대체합니다")
        create_pdf_simple()

    # DOCX 생성
    create_docx()

    # XLSX 생성
    create_xlsx()

    print()
    print("=" * 50)
    print(" 완료! demo_docs/ 폴더의 파일을 업로드하세요")
    print()
    print(" 시연 질문 예시:")
    print("   • 연차휴가는 몇 일인가요?")
    print("   • 초과근무수당은 어떻게 계산하나요?")
    print("   • 복리후생 제도를 알려주세요")
    print("   • 2025년 총 매출은 얼마인가요?")
    print("   • AI연구소 인원은 몇 명인가요?")
    print("   • BAIKAL Private AI 사용 방법을 알려주세요")
    print("=" * 50)
